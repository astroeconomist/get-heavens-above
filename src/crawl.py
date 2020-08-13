from bs4 import BeautifulSoup
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from PIL import Image
from time import sleep
from astro_calc import sun_h
from datetime import datetime
from os import makedirs, remove
import requests
import sys
from docx import Document
from docx.shared import Inches

LAT = 39.991489 #纬度
LNG = 116.308123 #经度
LOC = 'PKUJingyuan' #地点随便写
ALT = 50 #海拔
TZ = 'ChST' #时区
NUM = 3 #需要获取详细信息的过境事件个数


def parse(html):    #分析html文件，获取各过境信息
    global info
    soup = BeautifulSoup(html, features='lxml')
    try:
        table = soup.find('table', class_ = 'standardTable')
        for tr in table.tbody.find_all('tr'):
            line = []
            url = "https://heavens-above.com/{}".format(tr.td.a['href'])
            mjd = float(url.split('&mjd=')[1][:-7]) #从链接中获取ISS上中天时的mjd
            line.append(url)
            line.append(mjd)
            for td in tr.find_all('td'):
                line.append(td.string)
            info.append(line)
            #line的格式为[0链接,1 mjd, 2日期, 3星等, 4开始时间, 5高度角, 6方位, 7中天时间, 8中天高度角, 9中天方位, 10结束时间]
    except AttributeError: #本页没有内容
        return

def compare(event): #综合星等、太阳高度、时间、卫星中天高度计算得分，满分15分
    #计算ISS上中天时太阳高度角
    jd = event[1] + 2400000.5
    transit_sun_h = sun_h(jd, LNG, LAT)
    #计算太阳高度角得分，权重为30%
    if transit_sun_h < -18:
        sun_score = 10
    elif transit_sun_h < -12:
        sun_score = 10 - (transit_sun_h + 18) / 3
    elif transit_sun_h < -3:
        sun_score = 8 - (transit_sun_h + 12) / 9 * 8
    else:
        sun_score = 0
    #获取ISS中天星等
    mag = float(event[3])
    #计算星等得分，权重为40%
    mag_score = -10 * mag / 4
    #获取出现、消失、中天时间
    start_time = datetime.strptime(event[4], "%H:%M:%S")
    transit_time = datetime.strptime(event[7], "%H:%M:%S")
    end_time = datetime.strptime(event[10], "%H:%M:%S")
    #计算持续时间秒数
    period = (end_time - start_time).seconds
    #计算持续时间得分，权重为20%
    if period > 300:
        period_score = 10
    elif period > 200:
        period_score = 8
    elif period >100:
        period_score = 5
    else:
        period_score = 0
    #获取中天的高度角
    transit_h = float(event[8][:-1])
    #计算高度角得分，权重为10%
    transit_h_score = (transit_h - 10) / 8
    #前半夜的过境加5分
    is_before_midnight = 5 if transit_time.hour > 12 else 0
    #计算总分，满分为10+5=15分
    score = 0.4 * mag_score + 0.3 * sun_score + 0.2 * period_score + 0.1 * transit_h_score + is_before_midnight
    return score

def get_screenshot(driver, out):    #获取页面截图
    #利用js获取页面宽高
    width = driver.execute_script("return document.documentElement.scrollWidth")
    height = driver.execute_script("return document.documentElement.scrollHeight")
    #将浏览器的宽高设置成刚刚获取的宽高
    driver.set_window_size(width, height)
    sleep(1)
    #截图
    driver.save_screenshot("screen.png")
    #获取过境信息表格所在位置
    ele = driver.find_element_by_class_name("standardTable")
    left = ele.location['x']
    top = ele.location['y']
    right = left + ele.size['width']
    bottom = top + ele.size['height']
    im = Image.open('screen.png')
    im = im.crop((left, top, right, bottom))    #元素裁剪
    im.save(out)    #元素截图


def get_detail(driver, event, num): #下载星图并获取信息的截图
    url = event[0]
    driver.get(url)
    sleep(5)    #等一段时间让页面加载完毕
    html = driver.page_source
    soup = BeautifulSoup(html, features='lxml')
    #下载星图
    image_url = "https://heavens-above.com/{}".format(soup.find(id="ctl00_cph1_imgViewFinder").get('src'))
    r = requests.get(image_url, stream=True)
    with open('./out/img/img_{}.png'.format(num), 'wb') as f:
        for chunk in r.iter_content(chunk_size=32):
            f.write(chunk)
    #截图获取详细信息表格
    get_screenshot(driver, "./out/img/info_{}.png".format(num))

def write_docx(document, event, num):   #写《天象预报》的文件
    document.add_heading(event[2],level = 1)
    document.add_paragraph(u'{} {}等'.format(event[7], event[3]))
    document.add_picture('./out/img/img_{}.png'.format(num))
    document.add_picture('./out/img/info_{}.png'.format(num))
    document.add_paragraph('\n')

def main():
    makedirs("./out/img/", exist_ok=True)
    #设置虚拟浏览器
    chrome_options = Options()
    chrome_options.add_argument('headless')
    driver = webdriver.Chrome(chrome_options=chrome_options)
    #浏览主页面
    print('Getting the main page...')
    driver.get("https://heavens-above.com/PassSummary.aspx?satid=25544&lat={}&lng={}&loc={}&alt={}&tz={}".format(
        LAT, LNG, LOC, ALT, TZ))
    for _ in range(3):  #每页面有10天的信息，点击3次获取一个月的信息
        html = driver.page_source #Get the html source code.
        parse(html)
        sleep(1)
        driver.find_element_by_id("ctl00_cph1_btnNext").click()
    print('Done.')
    #按照评分标准进行排序
    sorted_info = sorted(info, key=compare, reverse=True)
    #新建word文档对象
    document = Document()
    document.add_heading(u'人造天体过境预报',0)
    for i in range(NUM):
        print("Downloading picture {} and screenshootting...".format(i+1))
        get_detail(driver, sorted_info[i], i)   #下载星图并获取信息的截图
        write_docx(document, sorted_info[i], i) #将其写入word文档
    driver.close()
    document.save("./out/output.docx")
    remove("screen.png")
    print("Done!")
    

if __name__ == "__main__":
    info = []
    main()

    

